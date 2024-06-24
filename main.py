import matplotlib.pyplot as plt
from docx import Document
import pandas as pd
import os, PIL.Image
from bs4 import BeautifulSoup
from markdown import markdown
from config import md_filename, docx_filename, gpt_model, gemini_model, gemini_vision, csv_path
from utils import read_data, generate_chart, generate_report, temperature, encode_image, save_as_docx, save_chart_query, read_prompts_from_yaml, pandas_agent_func, generate_report_gemini, generate_report_gpt4o
import json 
from docx.shared import Inches


def additional_prompt(data):
    additional_prompt = f"""
        Use the following data summaries for the report:
        Data Summary:
        {json.loads(data.describe(include="all").to_json())}
        """
    return additional_prompt

def add_graph_context(paragraph, image, explanation):
    p = paragraph.insert_paragraph_before()
    r = p.add_run()
    r.add_break()
    r.add_break()
    r.add_picture(image, width=Inches(6))  # Adjust width as needed
    r.add_break()
    r.add_break()
    r.add_break()
    r.add_text(explanation)
    r.add_break()

    return r

def main(data, report_pmt, chart_pmt, plot_model, use_model, use_vision_model, heading):
    if chart_pmt != []:
        for pmt in chart_pmt:
            print(pmt)
            updated_prompt = pmt + '\n Make sure to generate suitable charts to avoid any unneccessary errors. If any errors occur during the plot creation process, ignore them and do not produce any plots.'
            generate_chart(plot_model, data, updated_prompt)
    
    if report_pmt:
        print(report_pmt)
        addprompt = additional_prompt(data)
        report_pmt = addprompt+"\n\n"+report_pmt
        # print(report_pmt)
        # print("------------------------")
        result = generate_report(use_model, report_pmt,temperature=temperature)
        # print(result)

        new_prompt = result+'\n\n'+ 'This report was generated using AI, make it more relevant to business users and more insightful but tone down the statistic nature of the report. Make sure the main headings are unique and do not have same name in sub headings.'

        new_result= generate_report_gpt4o('user', new_prompt, num_token=2000, temperature=temperature)
        print(new_result)
        print("-----------------------------------")


    # Save the report to a md file
        with open(md_filename, "w", encoding="utf-8") as file:
            file.write(new_result)
        save_as_docx(md_filename, docx_filename)

    document = Document(docx_filename)

    imgs = [os.path.join("img",str(i)) for i in os.listdir("D:\\learning_project\\plot_pandasai\\img") if i.endswith(".png")]
    for paragraph in document.paragraphs:
        if heading in paragraph.text:
            print(paragraph.text)
            print("------------------------------------")
            for img in imgs:
                prompt = "Write a concise description of the image in less than 150 words. Just focus on min, max and new additions"
                # Add a new paragraph for each chart
                if use_vision_model == "gpt":
                    base64_image = encode_image(img)
                    content = [
                            {
                            "type": "text",
                            "text": prompt
                            },
                            {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                            }
                        ]
                    context = generate_report_gpt4o("user", content, 100)
                    add_graph_context(paragraph, img, context)
                elif use_vision_model ==  "gemini":
                    prompt = [prompt, PIL.Image.open(img)]
                    context = generate_report_gemini(gemini_vision, prompt)
                    add_graph_context(paragraph, img, context)

            break

    document.save("Final_Report_with_plot.docx") 

    if imgs != []:
        for file in imgs:
            file_path = os.path.join(file)
            os.remove(file_path)    


    return "Successfully Processed" 

if __name__ == "__main__":
    use_model = 'gpt' #'gemini'
    use_vision_model = "gpt"
    plot_model = "pandasai"
    heading = "Recommendations"

    input_df = read_data(csv_path)
    
    input_df.columns = ['index', 'order_id', 'date', 'status', 'fulfilment', 'sales_channel',
       'ship_service_level', 'style', 'sku', 'category', 'size', 'ASIN',
       'courier_status', 'qty', 'currency', 'amount', 'ship_city',
       'ship_state', 'ship_postal_code', 'ship_country', 'promotion-ids',
       'B2B', 'fulfilled_by', 'Unnamed:22']
    
    input_df['courier_status'] = input_df['courier_status'].fillna("NA")

    prompts = read_prompts_from_yaml("prompt.yaml")   
    charts = prompts[1]["charts"]
    report = prompts[0]["report"]
    main(input_df, report, charts, plot_model, use_model, use_vision_model, heading)

